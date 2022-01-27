# Code/logic adapted from Professor Clara James, Information Technology 1150: Programming and Logic / Al Sweigart,
# Automate the Boring Stuff With Python.
"""
This program will fetch and arrange data from the Minnesota State Park API server,
format and save park info in a Word document and create a travel guide.
Alternatively produces a sample guide if network or server connection is unavailable.
"""
import requests
import docx
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import random
import plotly.graph_objects as go
import plotly.express as px

document = docx.Document()


def main():
    """Outline data structure and save final document. """
    document.add_paragraph('Minnesota State Park Travel Guide', 'Title')

    response = master_park_data()
    park_dict = {}
    for mn_state_park in response:
        park_names_key = mn_state_park['name']
        park_id_value = mn_state_park['park_id']
        park_dict[park_names_key] = park_id_value
    park_ids_list = list(park_dict.values())

    number_of_parks = 5  # Guide should contain at least 5 parks
    random_park_id_selection = choose_rand_park(number_of_parks, park_ids_list)
    for park_id in random_park_id_selection:
        each_park = detailed_park_data(park_id)
        park_name = each_park['name']
        park_images = each_park['park_images']
        download_images(park_images, park_name)

        park_title_and_header_img(each_park)
        park_text_info(each_park)
        park_gallery(each_park)
        contact_information(each_park)
        open_street_park_map(each_park)
        document.add_paragraph()

    document.save('Minnesota_State_Park_Travel_Guide_Final.docx')


def master_park_data():
    """Attempt connection to the server and request master list of parks
    and their IDs. """
    api_mn_state_parks = 'https://mn-state-parks.herokuapp.com/api/list'
    try:
        data = requests.get(api_mn_state_parks).json()
        return data
    except:
        print('There was an error requesting park data. Check network connection.')


def detailed_park_data(park_id):
    """Request detailed information for each park. """
    api_park_detail = f'https://mn-state-parks.herokuapp.com/api/{park_id}'
    try:
        data = requests.get(api_park_detail).json()
        return data
    except:
        print('There was an error requesting park information. Check network '
              'connection.')


def choose_rand_park(total_parks, park_ids_list):
    """Use random module to choose parks from park list, return five random park IDs. """
    random_park_id_selection = []
    while total_parks > 0:
        chosen_park = random.choice(park_ids_list)
        if chosen_park not in random_park_id_selection:
            random_park_id_selection.append(chosen_park)
            total_parks -= 1

    return random_park_id_selection


def park_title_and_header_img(park_details):
    """Add park name and image header to document. """
    park_name = park_details['name']
    document.add_paragraph(park_name, 'Heading 1')
    document.add_picture(f'{park_name}_0.jpg', width=docx.shared.Inches(6),
                         height=docx.shared.Inches(2.49))


def park_text_info(park_details):
    """Create main body of text in the document for given park IDs. """

    highlights = park_details['highlights']
    document.add_paragraph('Highlights', 'Heading 2')
    for highlight in highlights:
        document.add_paragraph(highlight, 'List Bullet')

    park_information = park_details['park_information']
    for information_category, category_details in park_information.items():
        document.add_paragraph(f'{information_category}', 'Heading 2')
        document.add_paragraph(f'{category_details}', 'Normal')


def park_gallery(park_details):
    """Add remaining images to the document. """
    park_name = park_details['name']

    for index in range(6):  # Max 6 additional images per park in guide
        document.add_picture(f'{park_name}_{index+1}.jpg',
                             width=docx.shared.Inches(5.47),
                             height=docx.shared.Inches(2.87))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Code to center align image in docx
        # https://stackoverflow.com/questions/26474551/python-docx-lib-center-align-image


def open_street_park_map(park_details):
    """Request tile map using Mapbox open street map - does not require access token.
     Downloads image and adds map to travel guide. """
    park_name = park_details['name']
    lat = park_details['location']['latitude']
    lon = park_details['location']['longitude']

    fig = px.scatter_mapbox(
        lat=[lat],
        lon=[lon],
        color_discrete_sequence=["darkviolet"],
        zoom=4.1,
        width=420,
        height=300
    )
    fig.update_traces(marker=dict(size=9),
                      selector=dict(mode='markers'))
    fig.update_layout(mapbox_style="open-street-map")
    fig.update_layout(margin={"r": 0, "t": 0, "l": 0, "b": 0})
    fig.write_image(f'{park_name}_map.png')

    document.add_paragraph('Map', 'Heading 2')
    document.add_picture(f'{park_name}_map.png')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def contact_information(park_details):
    """Add contact information text to the document. """
    document.add_paragraph('Contact Information', 'Heading 2')

    document.add_paragraph('Address', 'Heading 3')
    document.add_paragraph(park_details['address'], 'Normal')

    document.add_paragraph('Website', 'Heading 3')
    document.add_paragraph(park_details['url'], 'Normal')


def download_images(park_images, park_name):
    """Take URLs from park dictionary and request, download, save and index images
    from MN park API server."""
    for index, url in enumerate(park_images):
        chosen_img = requests.get(url)
        filename = f'{park_name}_{index}.jpg'
        with open(filename, 'wb') as file:
            for chunk in chosen_img.iter_content():
                file.write(chunk)


main()
